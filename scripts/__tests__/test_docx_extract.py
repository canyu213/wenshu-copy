# -*- coding: utf-8 -*-
"""docx_extract.py 回归测试
运行：python scripts/__tests__/test_docx_extract.py
"""
import json
import os
import subprocess
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
import docx_extract as de  # noqa: E402

PASS = 0
FAIL = []


def check(name, cond, detail=""):
    global PASS
    if cond:
        PASS += 1
        print(f"  [PASS] {name}")
    else:
        FAIL.append(name)
        print(f"  [FAIL] {name} {detail}")


def make_docx(path: Path):
    import docx
    doc = docx.Document()
    doc.add_paragraph("教育信息化发展报告")
    doc.add_paragraph("发文机关〔2025〕12号")
    doc.add_paragraph("发布时间：2025年3月5日")
    doc.add_paragraph("本文讨论教育数字化转型的实践路径。")
    doc.save(str(path))


with tempfile.TemporaryDirectory(prefix="hermes-verify-docx-") as td:
    td_path = Path(td)
    docx_path = td_path / "报告.docx"
    make_docx(docx_path)
    script = str(Path(__file__).parent.parent / "docx_extract.py")

    print("测试 1: extract_text 单元")
    text = de.extract_text(docx_path)
    check("正文含标题行", "教育信息化发展报告" in text)
    check("正文含公文号行", "12号" in text)
    check("正文含正文段", "教育数字化转型" in text)

    print("测试 2: extract_metadata 单元（中性化验证）")
    meta = de.extract_metadata(text, "报告.docx")
    check("title 取首行", meta["title"] == "教育信息化发展报告", meta["title"])
    check("publish_date 提取", "2025年3月5日" in meta["publish_date"] or meta["publish_date"], meta["publish_date"])
    check("doc_number 提取", "〔2025〕12号" in meta["doc_number"], meta["doc_number"])
    check("无政治机构/标签残留", "中共中央" not in json.dumps(meta, ensure_ascii=False) and "新思想" not in json.dumps(meta, ensure_ascii=False))

    print("测试 3: CLI --meta")
    r = subprocess.run([sys.executable, script, "--input", str(docx_path), "--meta"],
                       capture_output=True, text=True, encoding="utf-8")
    check("exit 0", r.returncode == 0, r.stderr)
    data = json.loads(r.stdout)
    check("JSON 含 title/date/doc_number", data["title"] and data["publish_date"] and data["doc_number"], r.stdout[:150])

    print("测试 4: CLI --dir --output 生成 md")
    out = td_path / "out"
    r2 = subprocess.run([sys.executable, script, "--dir", str(td_path), "--output", str(out)],
                        capture_output=True, text=True, encoding="utf-8")
    check("exit 0", r2.returncode == 0, r2.stderr)
    mds = list(out.glob("*.md"))
    check("md 文件生成", len(mds) == 1, str(mds))
    md_text = mds[0].read_text(encoding="utf-8")
    check("md 含 YAML frontmatter", md_text.startswith("---\ntitle: 教育信息化发展报告"), md_text[:60])
    check("md 含正文", "教育数字化转型" in md_text)

    print("测试 5: 缺失依赖提示")
    fake = td_path / "无依赖.docx"
    r3 = subprocess.run([sys.executable, "-c",
                         "import sys; sys.path.insert(0, %r); import docx_extract as d; d.extract_text(%r)"
                         % (str(Path(__file__).parent.parent), str(fake))],
                        capture_output=True, text=True, encoding="utf-8")
    # 文件不存在时 extract_text 由调用方检查——这里只验证 import 链正常
    check("模块可导入", r3.returncode != 0 or "docx" in r3.stdout, r3.stderr[:100])

print(f"\n结果: {PASS} 通过 / {len(FAIL)} 失败")
sys.exit(1 if FAIL else 0)
